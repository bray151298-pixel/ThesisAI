"""Exportacion a Word (.docx) con formato oficial UNAJMA (Anexo 08).

Convierte el texto en Markdown que produce la IA (titulos ##, negritas **,
listas -, tablas | |) a formato Word real, y limpia restos de LaTeX.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------- limpieza LaTeX
_LATEX_SUBS = [
    (r"\$\$", ""), (r"\$", ""),
    (r"\\mathbf\{([^}]*)\}", r"\1"),
    (r"\\mathrm\{([^}]*)\}", r"\1"),
    (r"\\text\{([^}]*)\}", r"\1"),
    (r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1)/(\2)"),
    (r"\\sqrt\{([^}]*)\}", r"√(\1)"),
    (r"\\times", "×"), (r"\\cdot", "·"), (r"\\pm", "±"),
    (r"\\leq", "≤"), (r"\\geq", "≥"), (r"\\neq", "≠"), (r"\\approx", "≈"),
    (r"\\alpha", "α"), (r"\\beta", "β"), (r"\\sigma", "σ"), (r"\\mu", "μ"),
    (r"\\rho", "ρ"), (r"\\chi", "χ"), (r"\\sum", "Σ"),
    (r"\\left", ""), (r"\\right", ""),
    (r"\\begin\{[^}]*\}", ""), (r"\\end\{[^}]*\}", ""),
    (r"\\\\", " "),
    (r"\\[a-zA-Z]+", ""),   # cualquier comando LaTeX restante
]


def _clean_inline(text: str) -> str:
    for pat, rep in _LATEX_SUBS:
        text = re.sub(pat, rep, text)
    text = text.replace("`", "")
    # quita llaves sueltas de LaTeX pero conserva las de texto normal raras
    text = re.sub(r"[{}]", "", text)
    # HTML: <br> -> salto de linea; quita otras etiquetas comunes (sin tocar < > matematicos)
    text = re.sub(r"<\s*br\s*/?\s*>", "\n", text, flags=re.I)
    text = re.sub(
        r"</?\s*(b|i|u|p|strong|em|span|div|ul|ol|li|table|tr|td|th|h[1-6])\b[^>]*>",
        "", text, flags=re.I,
    )
    return text.strip()


def _norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (s or "").lower())


# ---------------------------------------------------------------- runs con negrita
def _add_runs(paragraph, text: str) -> None:
    text = _clean_inline(text)
    # cada linea (separada por <br> convertidos) se agrega con salto de linea real
    for li, line in enumerate(text.split("\n")):
        if li > 0:
            paragraph.add_run().add_break()
        for part in re.split(r"(\*\*[^*]+\*\*)", line):
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part:
                paragraph.add_run(part)


def _style_heading(heading) -> None:
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)


# ---------------------------------------------------------------- tablas markdown
def _is_sep(line: str) -> bool:
    s = line.strip().strip("|").replace(" ", "")
    return bool(s) and set(s) <= set("-:|")


def _add_table(doc: Document, rows_md: list[str]) -> None:
    rows = []
    for r in rows_md:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    table.allow_autofit = True
    for ri, cells in enumerate(rows):
        cells = cells + [""] * (ncols - len(cells))
        wrow = table.add_row().cells
        for ci, val in enumerate(cells):
            wrow[ci].text = ""
            p = wrow[ci].paragraphs[0]
            _add_runs(p, val)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
                if ri == 0:
                    run.bold = True


# ---------------------------------------------------------------- render de contenido
def _render_content(doc: Document, content: str, chapter_title: str) -> None:
    lines = (content or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # --- tabla markdown (fila con | seguida de separador ---) ---
        if "|" in stripped and i + 1 < len(lines) and _is_sep(lines[i + 1]):
            block = [line]
            i += 1  # separador (se descarta)
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                block.append(lines[i])
                i += 1
            _add_table(doc, block)
            continue

        # --- titulo markdown (#, ##, ###) ---
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            text = _clean_inline(m.group(2))
            if _norm(text) == _norm(chapter_title):  # evita duplicar el titulo del capitulo
                i += 1
                continue
            level = len(m.group(1))
            wl = 2 if level <= 2 else min(level, 4)
            h = doc.add_heading(text, level=wl)
            _style_heading(h)
            i += 1
            continue

        # --- lista con vinetas (- o *) ---
        mb = re.match(r"^[-*]\s+(.*)", stripped)
        if mb and not stripped.startswith("**"):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, mb.group(1))
            i += 1
            continue

        # --- recuadro de nota (> ...) ---
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            _add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # --- parrafo normal ---
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_runs(p, stripped)
        i += 1


# ---------------------------------------------------------------- numero de pagina
def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)


def _apply_unajma_format(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(4)
        section.top_margin = Cm(4)
        section.right_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_page_number(fp)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)


# ---------------------------------------------------------------- API principal
def build_docx(
    title: str,
    author: str,
    sections: list[tuple[str, str]],
    references: list[str],
) -> bytes:
    """Construye el .docx (formato UNAJMA) y devuelve los bytes para descargar."""
    doc = Document()
    _apply_unajma_format(doc)

    # ---- Portada ----
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run((title or "Tesis").upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(author)
    doc.add_page_break()

    # ---- Secciones (con render de Markdown) ----
    for heading, content in sections:
        h = doc.add_heading(_clean_inline(heading), level=1)
        _style_heading(h)
        _render_content(doc, content, heading)

    # ---- Referencias ----
    if references:
        doc.add_page_break()
        h = doc.add_heading("Referencias bibliográficas", level=1)
        _style_heading(h)
        for ref in references:
            p = doc.add_paragraph(ref)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
